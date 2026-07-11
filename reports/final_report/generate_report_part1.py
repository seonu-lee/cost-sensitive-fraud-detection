"""
비용 민감형 금융 사기탐지 모델 개발 및 통계적 성능 검증
1단계 최종보고서 생성 스크립트 (Part 1: 챕터 1~5)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 스타일 설정 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(6)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return h

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, bold=bold)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D6E4F0')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # 데이터
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(9)
            if c_idx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 열 너비
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ══════════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('비용 민감형 금융 사기탐지 모델 개발 및\n통계적 성능 검증')
run.font.name = '맑은 고딕'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('1단계 최종 보고서')
run2.font.name = '맑은 고딕'
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info_p.add_run('데이터: IEEE-CIS Fraud Detection (Kaggle)\n작성일: 2026년 7월')
run3.font.name = '맑은 고딕'
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 목차 (수동)
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '목차', 1)
toc_items = [
    ('1. 프로젝트 개요', ''),
    ('   1.1 프로젝트 목적', ''),
    ('   1.2 데이터셋 개요', ''),
    ('   1.3 프로젝트 차별점', ''),
    ('   1.4 전체 분석 흐름', ''),
    ('2. 데이터 탐색 및 통계검정', ''),
    ('   2.1 데이터 로딩 및 메모리 최적화', ''),
    ('   2.2 기초 탐색 및 클래스 분포', ''),
    ('   2.3 범주형 변수 통계검정', ''),
    ('   2.4 연속형 변수 통계검정', ''),
    ('   2.5 클래스 불균형 처리 결정', ''),
    ('3. Feature Engineering', ''),
    ('   3.1 UID 설계 및 사용자 단위 집계', ''),
    ('   3.2 금액 편차 파생변수', ''),
    ('   3.3 V컬럼 차원 축소 (PCA)', ''),
    ('   3.4 다중공선성 처리 및 최종 변수 세트', ''),
    ('4. Cross-Validation 전략 및 모델 비교', ''),
    ('   4.1 GroupKFold에서 StratifiedGroupKFold로', ''),
    ('   4.2 OOF(Out-of-Fold) 방식 채택', ''),
    ('   4.3 전처리 전략 확정', ''),
    ('   4.4 4개 모델 비교 결과', ''),
    ('5. 통계적 검증', ''),
    ('   5.1 DeLong Test (ROC-AUC 유의성)', ''),
    ('   5.2 Bootstrap CI (PR-AUC 불확실성)', ''),
    ('   5.3 McNemar Test (예측 패턴 차이)', ''),
    ('   5.4 Calibration 분석 및 개선 방향', ''),
    ('6. 하이퍼파라미터 튜닝', ''),
    ('   6.1 LightGBM: Optuna 튜닝', ''),
    ('   6.2 Random Forest: OOB 기반 탐색', ''),
    ('   6.3 튜닝 결과 및 최종 후보 선정', ''),
    ('7. Calibration 및 Threshold Optimization', ''),
    ('   7.1 확률 보정 결과 및 최종 모델 선택', ''),
    ('   7.2 Threshold Optimization (비용 최소화)', ''),
    ('8. 설명가능성 분석 (SHAP)', ''),
    ('   8.1 분석 모델 구성', ''),
    ('   8.2 전역 중요도 분석', ''),
    ('   8.3 개별 사례 해석', ''),
    ('9. Drift 분석 및 방법론적 자기 검증', ''),
    ('   9.1 분할 설계 및 목적 구분', ''),
    ('   9.2 연속형 변수 PSI / KS Test', ''),
    ('   9.3 범주형 변수 PSI', ''),
    ('   9.4 고카디널리티 PSI 착시 발견 및 재검증', ''),
    ('   9.5 시간적 누수 발견 및 보조 모델 재검증', ''),
    ('   9.6 card1 스무딩 강도 비교', ''),
    ('   9.7 종합 결론', ''),
    ('10. 종합 결론 및 실무적 시사점', ''),
    ('11. 한계 및 향후 과제', ''),
]
for item, _ in toc_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1장. 프로젝트 개요
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '1. 프로젝트 개요', 1)

add_heading(doc, '1.1 프로젝트 목적', 2)
add_para(doc, '본 프로젝트는 IEEE-CIS Fraud Detection 데이터셋(Kaggle)을 활용하여 금융 거래 사기탐지 모델을 개발하고, 그 성능과 안정성을 통계적으로 검증하는 것을 목적으로 함.')
add_para(doc, '사기탐지 도메인에서 모델 성능은 단순히 높은 AUC 수치만으로 평가될 수 없음. 사기 거래를 놓칠 때 발생하는 금전적 손실(FN 비용)과 정상 거래를 차단할 때 발생하는 고객 불만·운영 비용(FP 비용)은 비대칭적이며, 실무에서는 이 비용 구조를 모델에 명시적으로 반영하는 비용 민감형(cost-sensitive) 접근이 요구됨.')
add_para(doc, '본 프로젝트는 이러한 실무 요건에 대응하여 다음 세 가지 목표를 동시에 추구함.')

goals = [
    ('탐지 성능', '클래스 불균형 환경에서 실질적 사기탐지력(PR-AUC)을 극대화'),
    ('통계적 신뢰성', '모델 비교·선택의 모든 판단 근거를 통계 검정으로 정량화'),
    ('실무 적합성', '비용 구조에 맞춘 Threshold 최적화 및 확률 보정(Calibration)으로 운영 가능한 형태로 완성'),
]
add_table(doc,
    ['목표', '내용'],
    goals,
    col_widths=[3.5, 11.5])
add_caption(doc, '[표 1-1] 프로젝트 3대 목표')

add_heading(doc, '1.2 데이터셋 개요', 2)
add_para(doc, '본 데이터셋은 전자상거래 결제 플랫폼 Vesta Corporation이 제공한 실제 금융 거래 데이터로, 거래 정보(train_transaction)와 단말기·기기 정보(train_identity) 두 파일로 구성됨. 두 파일은 TransactionID를 기준으로 left join 방식으로 결합됨.')

add_table(doc,
    ['항목', '내용'],
    [
        ('전체 거래 건수', '590,540건'),
        ('수집 기간', '182일 (약 6.1개월)'),
        ('사기 거래 건수', '20,663건 (3.499%)'),
        ('정상 거래 건수', '569,877건 (96.501%)'),
        ('클래스 불균형 비율', '1 : 27.6'),
        ('원본 컬럼 수', '394개(거래) + 41개(기기) = 병합 후 435개'),
        ('기기 정보 보유 비율', '24.4% (144,233건)'),
    ],
    col_widths=[5, 10])
add_caption(doc, '[표 1-2] 데이터셋 기본 현황')

add_para(doc, '기기 정보 보유 거래의 사기 비율(7.85%)이 미보유 거래(2.09%)보다 3.7배 높아, 기기 정보의 존재 자체가 유의미한 신호임이 초기 탐색에서 확인됨.')

add_table(doc,
    ['변수군', '컬럼 수', '특징'],
    [
        ('V 컬럼', '339개', 'Vesta 독점 익명 변수, PCA 기반 구조 추정'),
        ('C 컬럼', '14개', '거래 관련 카운팅 변수'),
        ('D 컬럼', '15개', '시간 간격(time delta) 변수'),
        ('M 컬럼', '9개', '매칭 여부 플래그'),
        ('card 컬럼', '6개', '카드 정보 (card1: 고유값 13,553개)'),
        ('id 컬럼', '38개', '단말기·기기 식별 정보'),
        ('기타', '9개', 'addr, dist, email domain, ProductCD 등'),
    ],
    col_widths=[3, 2, 10])
add_caption(doc, '[표 1-3] 주요 변수군 구성')

add_heading(doc, '1.3 프로젝트 차별점', 2)
add_para(doc, '본 프로젝트의 핵심 차별점은 분석의 매 단계에 통계적 근거를 부여한다는 원칙임. 통계학·경영학 복수전공 배경을 살려, 단순한 모델 개발을 넘어 엄밀한 통계적 의사결정 체계를 구현하고자 함.')

diff_items = [
    ('통계적 EDA 및 변수 선별',
     '범주형: 카이제곱 검정 + Cramér\'s V / 연속형: 정규성 3종 검정 후 Mann-Whitney U + 효과크기(rank-biserial r) / 결측 패턴 자체를 독립 신호로 검증'),
    ('통계적 모델 비교',
     'DeLong Test(ROC-AUC), Bootstrap CI(PR-AUC), McNemar Test 병용 / Holm 다중비교 보정으로 6개 모델 쌍 동시 검증 시 제1종 오류 통제'),
    ('비용 민감형 Threshold 최적화',
     'FN/FP 비용 비율을 명시적으로 모델링하여 18가지 시나리오 분석 / 단일 F1 기준이 아닌 실제 비용 최소화 기준으로 운영 임계값 결정'),
    ('방법론적 자기 검증',
     'Day 9에서 기존 모델의 타겟 인코딩 방식이 시간적 Drift 검증과 상충함을 실증 발견 후 보조 모델로 자체 재검증 / 고카디널리티 PSI 왜곡(card1: 2.426→0.003) 등 분석 착시 직접 발견 및 보정'),
]
add_table(doc,
    ['차별점', '내용'],
    diff_items,
    col_widths=[4, 11])
add_caption(doc, '[표 1-4] 프로젝트 4대 차별점')

add_heading(doc, '1.4 전체 분석 흐름', 2)
add_para(doc, '본 프로젝트는 Day 01부터 Day 09까지 9단계로 구성되며, 각 단계의 산출물이 다음 단계의 입력으로 이어지는 순차적 파이프라인 구조를 가짐.')

add_table(doc,
    ['Day', '단계', '주요 내용'],
    [
        ('01', '데이터 로딩 및 기초 탐색', 'dtype 최적화 로딩, 클래스 분포 확인, V컬럼 그룹화'),
        ('02', '통계 검정', '카이제곱, Mann-Whitney U, 정규성 검정, 불균형 처리 결정'),
        ('03', 'Feature Engineering', 'UID 설계, 파생변수 생성, V컬럼 PCA, 다중공선성 처리'),
        ('04', 'CV 전략 및 모델 비교', 'StratifiedGroupKFold, OOF 방식, 4개 모델 비교'),
        ('05', '통계적 검증', 'DeLong Test, Bootstrap CI, McNemar Test, Holm 보정'),
        ('06', '하이퍼파라미터 튜닝', 'LightGBM Optuna 20 trials, RF OOB 탐색'),
        ('07', 'Calibration + Threshold', 'CalibratedClassifierCV(isotonic), 비용 최소화 Threshold'),
        ('08', '설명가능성 분석', 'SHAP TreeExplainer, 전역·개별 해석'),
        ('09', 'Drift 분석', 'PSI/KS Test, 고카디널리티 재그룹핑, 방법론적 자기 검증'),
    ],
    col_widths=[1.5, 4, 9.5])
add_caption(doc, '[표 1-5] 전체 분석 흐름')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2장. 데이터 탐색 및 통계검정
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '2. 데이터 탐색 및 통계검정', 1)
add_para(doc, '(Day 01 ~ Day 02에 해당)')

add_heading(doc, '2.1 데이터 로딩 및 메모리 최적화', 2)
add_para(doc, '분석 환경의 RAM 제약(16GB 미만)으로 인해 원시 데이터를 그대로 로딩하면 메모리 부족이 발생할 수 있음. 이를 해결하기 위해 각 컬럼의 실제 값 범위를 사전 파악한 후, 정보 손실 없이 dtype을 최소 크기로 다운캐스팅하는 방식을 채택함.')
add_para(doc, '다운캐스팅 결과, float32 최대 오차는 8.79e-05로 모델 학습에 무시 가능한 수준임이 검증됨(오차 발생 비율 34.4%, 단 절대값 극소).')

add_table(doc,
    ['파일', '원본 크기', '최적화 후', '감소율'],
    [
        ('train_transaction', '2,062.07 MB', '861.12 MB', '58.2%'),
        ('train_identity', '143.14 MB', '16.17 MB', '88.7%'),
        ('병합 데이터', '약 921 MB', '약 922 MB (결합 포함)', '—'),
    ],
    col_widths=[4, 3, 3, 2.5])
add_caption(doc, '[표 2-1] dtype 최적화에 의한 메모리 감소 효과')

add_heading(doc, '2.2 기초 탐색 및 클래스 분포', 2)
add_para(doc, '병합 후 데이터셋은 590,540행 × 435열(has_identity 포함) 구조임. 주요 탐색 결과는 다음과 같음.')

add_para(doc, '클래스 분포 및 불균형 현황', bold=True)
add_para(doc, '사기 거래는 전체의 3.499%(20,663건)에 불과하며, 정상 거래(569,877건) 대비 클래스 불균형 비율이 1:27.6에 달함. 이는 단순 정확도(Accuracy) 지표가 무의미하며, PR-AUC를 주 평가 지표로 채택해야 함을 의미함.')

add_para(doc, '기기 정보(has_identity)와 사기 비율', bold=True)
add_table(doc,
    ['구분', '거래 건수', '비율', '사기 비율'],
    [
        ('기기 정보 보유 (has_identity=1)', '144,233건', '24.4%', '7.85%'),
        ('기기 정보 미보유 (has_identity=0)', '446,307건', '75.6%', '2.09%'),
    ],
    col_widths=[5.5, 3, 2.5, 2.5])
add_caption(doc, '[표 2-2] 기기 정보 보유 여부에 따른 사기 비율 비교')
add_para(doc, '기기 정보 보유 거래의 사기 비율이 미보유 거래 대비 3.7배 높아, has_identity 플래그가 유의미한 신호임이 확인됨. 이는 Day 2의 카이제곱 검정으로 통계적으로 재검증됨.')

add_para(doc, 'V컬럼 결측 패턴 그룹화', bold=True)
add_para(doc, '339개의 V컬럼은 결측 여부의 이진 패턴 기준으로 15개의 그룹으로 완벽하게 분해됨(패턴 일치도 100%). 이는 V컬럼이 독립 변수들의 집합이 아닌, 그룹별로 함께 생성·사용되는 구조적 설계임을 시사함. 이 발견은 Day 3의 그룹별 PCA 전략의 근거가 됨.')

add_para(doc, '결측률 분포', bold=True)
add_table(doc,
    ['결측률 구간', '컬럼 수', '대표 컬럼'],
    [
        ('결측치 없음', '21개', 'C 컬럼 전체, card1 등'),
        ('0~1%', '112개', '주요 수치형 변수'),
        ('10~50%', '108개', 'D, addr 컬럼 일부'),
        ('50~90%', '202개', 'M 컬럼, V 컬럼 일부'),
        ('90% 이상', '12개', 'id_24(99.2%), id_25(99.1%) 등'),
    ],
    col_widths=[3.5, 2.5, 7.5])
add_caption(doc, '[표 2-3] 결측률 구간별 컬럼 수 분포')

add_heading(doc, '2.3 범주형 변수 통계검정', 2)
add_para(doc, '범주형 변수와 isFraud의 연관성을 카이제곱 독립성 검정으로 검증하고, 효과크기는 Cramér\'s V로 정량화함. 대표 변수 5개의 결과는 다음과 같음.')

add_table(doc,
    ['변수', 'χ² 통계량', 'p-value', 'Cramér\'s V', '유의'],
    [
        ('ProductCD', '16,742.17', '0.0', '0.1684', 'True'),
        ('has_identity', '10,683.64', '0.0', '0.1345', 'True'),
        ('card6', '5,957.03', '0.0', '0.1006', 'True'),
        ('DeviceType', '609.62', '1.35e-134', '0.0658', 'True'),
        ('card4', '364.87', '8.97e-79', '0.0249', 'True'),
    ],
    col_widths=[3.5, 3, 2.5, 2.5, 1.5])
add_caption(doc, '[표 2-4] 주요 범주형 변수 카이제곱 검정 결과 (유의수준 0.05)')

add_para(doc, '결측 플래그를 별도 신호로 검정한 결과, 결측 자체가 이진 변수로서 강한 연관성을 보임이 확인됨.')
add_table(doc,
    ['결측 플래그 변수', 'χ² 통계량', 'Cramér\'s V', '유의'],
    [
        ('addr_missing', '15,016.72', '0.1595', 'True'),
        ('has_identity', '10,683.64', '0.1345', 'True'),
        ('M123_missing', '4,720.58', '0.0894', 'True'),
        ('M789_missing', '2,876.27', '0.0698', 'True'),
    ],
    col_widths=[4.5, 3, 3, 2.5])
add_caption(doc, '[표 2-5] 결측 플래그 변수 카이제곱 검정 결과')

add_para(doc, 'V컬럼 15개 그룹의 결측 플래그도 전부 검정한 결과, 12개 그룹이 유의미한 연관성을 보였으며(p < 0.05), V_group1·2·3만 Cramér\'s V가 0.004 미만으로 제외됨. 이를 통해 V컬럼 결측 여부가 독립적 예측 신호임을 통계적으로 확인함.')

add_heading(doc, '2.4 연속형 변수 통계검정', 2)
add_para(doc, '연속형 변수의 사기 여부 간 분포 차이를 검정하기 전, 정규성 여부를 3단계로 확인함.')

add_para(doc, '정규성 검정 (3종 병용)', bold=True)
add_para(doc, '① D\'Agostino-Pearson Test (왜도+첨도 기반, 5만 건 샘플링, 주 검정) / ② Anderson-Darling Test (분포 꼬리 차이에 민감, 보강) / ③ Shapiro-Wilk Test (5,000건 샘플, 참고용)')
add_para(doc, 'C1~C14, TransactionAmt, D1 등 16개 변수 전부 D\'Agostino-Pearson p=0.0으로 비정규 확인됨(예: C3 왜도=88.95, 첨도=11,000). 이에 따라 t-검정 대신 비모수 검정인 Mann-Whitney U 검정을 채택함.')

add_para(doc, 'Mann-Whitney U 검정 결과 (효과크기 기준 정렬)', bold=True)
add_table(doc,
    ['변수', '정상 중앙값', '사기 중앙값', '효과크기 r', 'p-value', '유의'],
    [
        ('C4', '0.0', '1.0', '0.3754', '0.0', 'True'),
        ('C8', '0.0', '1.0', '0.3664', '0.0', 'True'),
        ('C10', '0.0', '1.0', '0.3558', '0.0', 'True'),
        ('C12', '0.0', '0.0', '0.3269', '0.0', 'True'),
        ('C7', '0.0', '0.0', '0.2949', '0.0', 'True'),
        ('D1', '4.0', '0.0', '0.1899', '0.0', 'True'),
        ('amt_robust_zscore', '0.00', '0.25', '0.1157', '9.67e-155', 'True'),
        ('amt_zscore', '-0.29', '-0.16', '0.0993', '4.51e-125', 'True'),
        ('TransactionAmt', '68.5', '75.0', '0.0050', '0.226', 'False ※'),
    ],
    col_widths=[3.5, 2.5, 2.5, 2.5, 2.5, 1.5])
add_caption(doc, '[표 2-6] 주요 연속형 변수 Mann-Whitney U 검정 결과 (효과크기 기준 정렬)')

add_para(doc, '※ TransactionAmt 원본은 정상·사기 간 분포 차이가 통계적으로 유의하지 않음(p=0.226). 그러나 UID 그룹 내 상대적 편차로 변환한 amt_robust_zscore는 effect size 0.1157로 유의미한 신호로 전환됨. 이는 Day 3에서 금액 편차 파생변수를 설계한 근거임.')

add_heading(doc, '2.5 클래스 불균형 처리 결정', 2)
add_para(doc, '사기 거래 비율이 3.499%로 클래스 불균형이 심각하므로, 처리 방식을 통계적 근거에 따라 결정함.')
add_para(doc, '결정: scale_pos_weight(class_weight=\'balanced\') 채택 / SMOTE·언더샘플링 보류')

add_table(doc,
    ['후보 방법', '검토 결과'],
    [
        ('scale_pos_weight = 27.58', '채택 — LightGBM·XGBoost 기본 지원, 정보 손실 없음'),
        ('SMOTE 오버샘플링', '보류 — V컬럼이 PCA 기반 익명 변수이므로 보간의 의미 불확실'),
        ('언더샘플링', '보류 — C컬럼 등 유의미한 신호 정보를 버리는 비용 발생'),
    ],
    col_widths=[4.5, 10.5])
add_caption(doc, '[표 2-7] 클래스 불균형 처리 방식 비교 및 결정')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3장. Feature Engineering
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '3. Feature Engineering', 1)
add_para(doc, '(Day 03에 해당)')

add_heading(doc, '3.1 UID 설계 및 사용자 단위 집계', 2)
add_para(doc, 'card1 단독으로는 13,553개의 고유값을 가져 카드 번호 수준의 식별이 가능하나, 거래 주체(사용자)를 안정적으로 추적하려면 추가 정보와의 결합이 필요함. 후보 조합들을 비교한 결과 card1+card2+addr1 조합을 UID로 채택함.')

add_table(doc,
    ['UID 구성 후보', '고유값 수', '1건짜리 비율(그룹)', '채택 여부'],
    [
        ('card1만', '13,553개', '—', '미채택'),
        ('card1+card2만', '14,524개', '—', '미채택'),
        ('card1+addr1만', '39,974개', '—', '미채택'),
        ('card1+card2+addr1 (채택)', '41,672개', '40.3%', '채택'),
        ('card1+card2+addr1+D1', '248,038개', '73.4%', '미채택'),
    ],
    col_widths=[5, 2.5, 3.5, 2])
add_caption(doc, '[표 3-1] UID 후보 조합 비교')

add_para(doc, 'D1은 거래 시점마다 달라지는 시간 간격(timedelta) 변수이므로 그룹화에 부적합하여 제외함. 채택된 UID는 평균 14.17건/UID, 최대 9,900건의 거래를 묶으며, 이후 UID 단위 금액 편차 변수 산출의 기반이 됨.')

add_heading(doc, '3.2 금액 편차 파생변수', 2)
add_para(doc, 'Day 2에서 TransactionAmt 원본이 유의하지 않음을 확인(p=0.226)했으므로, 절대 금액이 아닌 사용자 그룹 내 상대적 편차를 포착하는 파생변수를 설계함.')

add_para(doc, 'Robust Z-score 채택 이유', bold=True)
add_para(doc, '표준 Z-score는 평균과 표준편차를 사용하므로 이상치에 민감함. 본 데이터의 거래 금액은 왜도가 14.4, 첨도 1,124로 극단적 분포이므로, median과 MAD(Median Absolute Deviation) 기반의 Robust Z-score를 채택함.')
add_para(doc, '공식: (x - median) / (1.4826 × MAD)  ※ 1.4826은 MAD를 표준편차와 동일 척도로 맞추기 위한 보정 상수')

add_para(doc, 'MAD floor 설정 (0.3 채택)', bold=True)
add_para(doc, 'UID 그룹 내에서 MAD=0인 그룹이 21,612개(단일 금액 반복 사용자)로, 분모가 0이 되면 Z-score가 무한대로 발산함. 이를 방지하기 위해 MAD 최솟값에 floor를 적용하고, 결측률·std를 기준으로 floor_ratio=0.3을 채택함.')

add_table(doc,
    ['floor_ratio', 'floor 값', '결측 비율', 'std'],
    [
        ('0.05', '1.94', '6.63%', '4.91'),
        ('0.10', '3.88', '7.19%', '4.73'),
        ('0.20', '7.76', '9.45%', '4.22'),
        ('0.30 (채택)', '11.64', '12.83%', '3.88'),
        ('0.50', '19.40', '26.50%', '3.92'),
    ],
    col_widths=[3.5, 2.5, 3, 2])
add_caption(doc, '[표 3-2] MAD floor 후보별 결측률 및 분포 안정성 비교')

add_table(doc,
    ['변수', 'effect_size_r', 'p-value', '채택 여부'],
    [
        ('amt_zscore (표준 Z-score)', '0.09930', '4.51e-125', '비교 기준'),
        ('amt_robust_zscore (채택)', '0.11573', '9.67e-155', '채택 (+16.6% 개선)'),
        ('amt_hour_robust_zscore', '0.00500', '0.205', '제외 (유의하지 않음)'),
    ],
    col_widths=[5, 2.5, 2.5, 3])
add_caption(doc, '[표 3-3] 금액 편차 파생변수 효과크기 비교')

add_para(doc, '시간대별 Z-score(amt_hour_robust_zscore)는 effect_size=0.005, p=0.205로 유의하지 않아 최종 변수 세트에서 제외함.')

add_heading(doc, '3.3 V컬럼 차원 축소 (PCA)', 2)
add_para(doc, 'V컬럼 339개를 모두 모델에 투입하면 다중공선성과 차원의 저주 문제가 발생함. Day 1에서 확인한 15개 결측 패턴 그룹을 기반으로 그룹별 PCA를 적용함. 그룹 내 변수들이 구조적으로 함께 생성된 것이므로, 그룹별 PCA는 의미론적으로도 타당한 압축 전략임.')

add_para(doc, '그룹별 목표 설명분산은 그룹 내 평균 절대 상관계수를 고려하여 차등 적용함(상관 높은 그룹 95%, 낮은 그룹 85%). 전체 압축 결과는 다음과 같음.')

add_table(doc,
    ['압축 대상', '원본 수', '압축 후', '설명분산'],
    [
        ('V_group1~15 (13개 그룹)', '317개', '89개 주성분', '85~96% (그룹별 상이)'),
        ('V_group13+14 통합', '22개', '11개 주성분', '91.7%'),
        ('C컬럼군 (11개)', '11개', '2개 주성분', '97.4%'),
        ('결측 플래그 (25개)', '25개', '9개 주성분 (AllFlag)', '91.3%'),
    ],
    col_widths=[4.5, 2, 2.5, 4])
add_caption(doc, '[표 3-4] PCA 압축 결과 요약')
add_para(doc, '전체적으로 V컬럼 339개가 89개 주성분으로 압축되어 약 72%의 차원 감소가 이루어짐.')

add_heading(doc, '3.4 다중공선성 처리 및 최종 변수 세트', 2)
add_para(doc, 'D9와 Transaction_hour의 상관계수=1.000, C7과 C12의 상관계수=0.999 등 121개의 0.8 이상 상관 쌍이 발견됨. VIF(Variance Inflation Factor) 검증 및 추가 PCA 통합 처리를 수행함.')

add_table(doc,
    ['처리 전후', 'VIF 최댓값', 'VIF 10 이상 비율'],
    [
        ('처리 전', '86,308 (V_group15_missing)', '—'),
        ('처리 후', '572 (V_group2_PC1)', '43개 / 157개 (27%)'),
    ],
    col_widths=[3.5, 5, 4.5])
add_caption(doc, '[표 3-5] VIF 처리 결과 (99.3% 개선)')

add_para(doc, '트리 기반 모델(LightGBM, RF, XGBoost)은 다중공선성에 강건하므로 잔존 VIF가 모델 성능에 미치는 영향은 제한적임. LR에는 L2 정규화를 적용함.')

add_table(doc,
    ['항목', '값'],
    [
        ('최종 데이터 크기', '590,540행 × 208열'),
        ('모델 입력 변수 수', '203개 (target 및 ID 제외)'),
        ('결측치', '0'),
        ('파일 크기', '105.87 MB (day3_final.parquet)'),
    ],
    col_widths=[5, 8])
add_caption(doc, '[표 3-6] 최종 변수 세트 요약')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4장. Cross-Validation 전략 및 모델 비교
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '4. Cross-Validation 전략 및 모델 비교', 1)
add_para(doc, '(Day 04에 해당)')

add_heading(doc, '4.1 GroupKFold에서 StratifiedGroupKFold로', 2)
add_para(doc, '사기탐지 모델에서 동일 카드 소유자의 거래가 학습과 검증 세트에 동시에 포함될 경우 데이터 누수가 발생할 수 있음. 이를 방지하기 위해 UID(card1+card2+addr1)를 그룹 기준으로 하는 그룹 기반 분할을 적용함.')
add_para(doc, '초기에 GroupKFold 5-fold를 적용했을 때, 평가 지표에 구조적 왜곡이 발견됨.')

add_table(doc,
    ['Fold', '거래 건수', '사기 비율', '비고'],
    [
        ('Fold 1', '118,108건', '3.94%', '—'),
        ('Fold 2', '118,108건', '2.87%', '※ 이상 저하'),
        ('Fold 3', '118,108건', '3.60%', '—'),
        ('Fold 4', '118,108건', '3.17%', '—'),
        ('Fold 5', '118,108건', '3.93%', '—'),
    ],
    col_widths=[2, 3, 3, 5])
add_caption(doc, '[표 4-1] GroupKFold 적용 시 fold별 사기 비율 (교체 전)')

add_para(doc, 'Fold 2의 사기 비율(2.87%)이 다른 fold(3.17~3.94%)보다 유독 낮아, 해당 fold의 PR-AUC가 왜곡되는 문제가 확인됨. 이는 단순히 랜덤 변동이 아닌, UID 그룹의 구조적 불균등 배치에서 기인한 문제임.')
add_para(doc, 'StratifiedGroupKFold로 교체한 결과, 그룹 경계를 유지하면서 동시에 클래스 비율을 균등화함.')

add_table(doc,
    ['Fold', '거래 건수', '사기 비율', '비고'],
    [
        ('Fold 1', '118,110건', '3.50%', '균등화'),
        ('Fold 2', '118,110건', '3.50%', '균등화'),
        ('Fold 3', '118,108건', '3.50%', '균등화'),
        ('Fold 4', '118,109건', '3.50%', '균등화'),
        ('Fold 5', '118,103건', '3.50%', '균등화'),
    ],
    col_widths=[2, 3, 3, 5])
add_caption(doc, '[표 4-2] StratifiedGroupKFold 적용 시 fold별 사기 비율 (교체 후)')

add_para(doc, '이 개선 과정은 단순히 라이브러리를 교체한 것이 아니라, 평가 지표 왜곡을 실증적으로 발견하고 방법론을 수정한 사례로서 의미를 가짐.')

add_heading(doc, '4.2 OOF(Out-of-Fold) 방식 채택', 2)
add_para(doc, '별도의 holdout test set을 구성하는 대신, StratifiedGroupKFold 5-fold OOF(Out-of-Fold) 방식을 채택함. 각 샘플의 예측값은 해당 샘플이 학습에 포함되지 않은 fold에서 생성되므로, 전체 590,540건에 대한 불편 예측이 확보됨. 특히 사기 거래(20,663건)처럼 소수 클래스 샘플이 귀한 경우, holdout으로 일부를 고정하는 방식보다 전체 샘플을 최대한 활용할 수 있는 OOF 방식이 유리함.')

add_heading(doc, '4.3 전처리 전략 확정', 2)
add_table(doc,
    ['항목', '채택 방법', '근거'],
    [
        ('수치형 스케일링', 'StandardScaler', 'MinMaxScaler는 극단값이 스케일 지배; RobustScaler는 파생변수에 이미 적용; L2 정규화와 수학적으로 잘 맞는 StandardScaler 선택'),
        ('범주형 인코딩 (LR/RF)', '타겟 인코딩 (smoothing=20, fold 내부)', '고카디널리티 변수(card1: 13,553개)의 one-hot 차원 폭발 방지'),
        ('범주형 인코딩 (LightGBM/XGBoost)', 'category 네이티브', '내장 처리로 추가 변환 불필요'),
    ],
    col_widths=[3.5, 3.5, 8])
add_caption(doc, '[표 4-3] 전처리 전략 결정 사항')

add_heading(doc, '4.4 4개 모델 비교 결과', 2)
add_para(doc, 'Logistic Regression(베이스라인), Random Forest, XGBoost, LightGBM 4개 모델을 동일 조건(StratifiedGroupKFold 5-fold OOF, 클래스 불균형 처리 통일)으로 비교함.')

add_table(doc,
    ['모델', 'PR-AUC (mean±std)', 'ROC-AUC (mean±std)', '순위'],
    [
        ('Random Forest', '0.4972 ± 0.0342', '0.8632 ± 0.0089', '1위'),
        ('LightGBM', '0.4860 ± 0.0304', '0.8463 ± 0.0114', '2위'),
        ('XGBoost', '0.4471 ± 0.0337', '0.8101 ± 0.0169', '3위'),
        ('Logistic Regression', '0.3366 ± 0.0261', '0.8235 ± 0.0182', '4위'),
    ],
    col_widths=[4, 4.5, 4.5, 2])
add_caption(doc, '[표 4-4] 4개 모델 OOF 성능 비교 결과 (PR-AUC 기준 정렬)')

add_para(doc, 'Random Forest가 PR-AUC와 ROC-AUC 양쪽에서 모두 1위를 기록함. LightGBM과의 PR-AUC 차이는 0.0112로, 이 차이의 통계적 유의성은 Day 5에서 검증됨.')
add_para(doc, 'XGBoost의 ROC-AUC(0.8101)가 Logistic Regression(0.8235)보다 낮은 역전 현상은 scale_pos_weight=27.6 설정이 소수 클래스 탐지를 과도하게 강조하여 전반적 구분력이 저하된 결과로 해석됨.')
add_para(doc, 'Fold 4의 PR-AUC가 전 모델에 걸쳐 동시에 저하되는 현상은, 1건짜리 UID(사용자 패턴 파악 불가)나 고빈도 UID(추정 식별자 신뢰도 낮음)가 Fold 4에 집중된 구조적 특성에 기인함으로 해석되며, UID가 완벽한 개인 식별자가 아니라는 데이터 구조적 한계와 일관된 결과임.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5장. 통계적 검증
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '5. 통계적 검증', 1)
add_para(doc, '(Day 05에 해당)')
add_para(doc, 'Day 4의 성능 수치가 실질적인 차이인지, 아니면 표집 변동에 불과한지를 3종 통계 검정으로 체계적으로 검증함. 검정에 사용된 예측값은 모두 OOF(Out-of-Fold) 예측으로, 590,540건 전체에 대한 불편 예측임.')

add_heading(doc, '5.1 DeLong Test (ROC-AUC 유의성)', 2)
add_para(doc, '두 모델의 ROC-AUC 차이가 통계적으로 유의한지 placement value 기반 DeLong 분산 추정으로 검정함. 4개 모델 간 6개 쌍에 대해 Holm 다중비교 보정을 적용하여 제1종 오류를 통제함.')

add_table(doc,
    ['모델 A', '모델 B', 'AUC_A', 'AUC_B', '차이', 'z 통계량', 'p (Holm)', '유의'],
    [
        ('LR', 'RF', '0.8229', '0.8632', '-0.0403', '-28.08', '0.0', 'True'),
        ('LR', 'XGBoost', '0.8229', '0.8099', '+0.0130', '7.06', '1.66e-12', 'True'),
        ('LR', 'LightGBM', '0.8229', '0.8465', '-0.0235', '-14.69', '0.0', 'True'),
        ('RF', 'XGBoost', '0.8632', '0.8099', '+0.0533', '31.80', '0.0', 'True'),
        ('RF', 'LightGBM', '0.8632', '0.8465', '+0.0167', '11.67', '0.0', 'True'),
        ('XGBoost', 'LightGBM', '0.8099', '0.8465', '-0.0366', '-36.70', '0.0', 'True'),
    ],
    col_widths=[2, 2.5, 1.7, 1.7, 1.5, 2.3, 2.3, 1])
add_caption(doc, '[표 5-1] DeLong Test 결과 (Holm 다중비교 보정 적용)')

add_para(doc, 'ROC-AUC 순위: RF(0.8632) > LightGBM(0.8465) > LR(0.8229) > XGBoost(0.8099). 6개 쌍 전부 통계적으로 유의함. 특히 RF와 LightGBM 간에도 z=11.67, p_holm≈0으로 유의한 차이가 확인됨. 다만 59만 건의 대표본에서는 작은 차이도 유의하게 나타날 수 있어, 효과 크기와 함께 해석할 필요가 있음.')

add_heading(doc, '5.2 Bootstrap CI (PR-AUC 불확실성)', 2)
add_para(doc, 'PR-AUC 차이의 신뢰 구간을 Bootstrap 1,000회 재표집으로 추정함. PR-AUC는 불균형 클래스 환경에서 ROC-AUC보다 실질적 탐지 성능을 더 직접적으로 반영하는 지표이므로 본 검증의 주 지표로 사용함.')

add_table(doc,
    ['모델 A', '모델 B', 'PR-AUC_A', 'PR-AUC_B', '차이', '95% CI', '유의'],
    [
        ('LR', 'RF', '0.3340', '0.4971', '-0.1631', '[-0.1680, -0.1576]', 'True'),
        ('LR', 'XGBoost', '0.3340', '0.4449', '-0.1109', '[-0.1161, -0.1054]', 'True'),
        ('LR', 'LightGBM', '0.3340', '0.4848', '-0.1508', '[-0.1558, -0.1454]', 'True'),
        ('RF', 'XGBoost', '0.4971', '0.4449', '+0.0522', '[+0.0477, +0.0563]', 'True'),
        ('RF', 'LightGBM', '0.4971', '0.4848', '+0.0122', '[+0.0089, +0.0158]', 'True ※'),
        ('XGBoost', 'LightGBM', '0.4449', '0.4848', '-0.0399', '[-0.0428, -0.0371]', 'True'),
    ],
    col_widths=[2, 2.5, 2, 2, 1.8, 3.7, 1])
add_caption(doc, '[표 5-2] Bootstrap CI 결과 (PR-AUC 차이, n_bootstrap=1,000, 95% CI)')

add_para(doc, '※ RF vs LightGBM: 차이 0.0122, 95% CI [0.0089, 0.0158]. CI가 0을 포함하지 않아 통계적으로 유의하며, 이 시점에서 RF가 LightGBM보다 유의하게 우세함이 확인됨. 다만 절대 차이가 0.0122로 작아, 하이퍼파라미터 튜닝 후 역전 가능성이 있음을 시사함.')

add_heading(doc, '5.3 McNemar Test (예측 패턴 차이)', 2)
add_para(doc, '동일 관측치에 대한 두 모델의 예측 일치·불일치 패턴을 카이제곱 기반 McNemar Test(연속성 보정 적용)로 검증함. threshold=0.5를 적용한 이진 예측 비교이므로 보조 확인용으로 활용함.')

add_table(doc,
    ['모델 A', '모델 B', 'A만 정답', 'B만 정답', 'χ²', 'p (Holm)', '유의'],
    [
        ('LR', 'RF', '7,808건', '69,503건', '49,231.67', '0.0', 'True'),
        ('LR', 'XGBoost', '7,900건', '66,056건', '45,729.95', '0.0', 'True'),
        ('LR', 'LightGBM', '11,868건', '58,614건', '31,002.17', '0.0', 'True'),
        ('RF', 'XGBoost', '7,238건', '3,699건', '1,144.50', '7.02e-251', 'True'),
        ('RF', 'LightGBM', '21,030건', '6,081건', '8,241.77', '0.0', 'True'),
        ('XGBoost', 'LightGBM', '15,845건', '4,435건', '6,418.41', '0.0', 'True'),
    ],
    col_widths=[2, 2.5, 2.5, 2.5, 2.5, 2.5, 1])
add_caption(doc, '[표 5-3] McNemar Test 결과 (Holm 보정, threshold=0.5)')

add_para(doc, 'RF vs LightGBM: RF만 맞힌 21,030건 vs LightGBM만 맞힌 6,081건(약 3.5배). RF가 LightGBM에 비해 훨씬 많은 케이스를 단독으로 정확히 분류하고 있으며, 이는 DeLong Test·Bootstrap CI 결과와 일관된 방향임.')

add_heading(doc, '5.4 Calibration 분석 및 개선 방향', 2)
add_para(doc, '사기탐지 모델에서 예측 확률은 Threshold 최적화의 기반이 되므로, 모델이 출력하는 확률값의 신뢰도(Calibration)가 중요함. 모델 예측 확률을 10개 구간으로 나누고 실제 사기 비율과의 평균 절대 차이(|Gap|)를 측정함.')

add_table(doc,
    ['모델', '평균 |Gap|', '최대 |Gap|', '방향', '개선 여지'],
    [
        ('Random Forest', '0.1366', '0.2560', '과소신뢰 (실제 > 예측)', '상대적으로 작음'),
        ('XGBoost', '0.2485', '0.4588', '과신뢰 (예측 > 실제)', '중간'),
        ('LightGBM', '0.3130', '0.5473', '과신뢰 (예측 > 실제)', '가장 큼 ※'),
        ('Logistic Regression', '0.4007', '0.6792', '과신뢰 (예측 > 실제)', '—'),
    ],
    col_widths=[3.5, 2.5, 2.5, 4, 2.5])
add_caption(doc, '[표 5-4] Calibration 분석 결과 (평균 |Gap| 기준)')

add_para(doc, '※ Day 5 시점의 종합 결론: 3개 검정 모두에서 RF가 LightGBM을 통계적으로 유의하게 앞섰으며, 이 시점에서는 RF가 우세한 상태임. 그러나 LightGBM의 Calibration Gap(0.3130)이 RF(0.1366) 대비 2.3배 크다는 점은, 확률 보정을 통해 LightGBM이 더 큰 개선 여지를 가진다는 것을 시사함. 이 가능성을 Day 6(하이퍼파라미터 튜닝)과 Day 7(Calibration)에서 실제로 검증함.')

doc.add_page_break()

# 저장 (Part 1)
import os
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fraud_detection_report_draft.docx')
doc.save(out_path)
print(f"Part 1 저장 완료: {out_path}")
