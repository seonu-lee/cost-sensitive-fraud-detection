"""
비용 민감형 금융 사기탐지 모델 개발 및 통계적 성능 검증
1단계 최종보고서 생성 스크립트 (Part 2: 챕터 6~11)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Part 1에서 저장한 파일 이어서 열기
draft_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fraud_detection_report_draft.docx')
doc = Document(draft_path)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return h

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, bold=bold)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D6E4F0')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(9)
            if c_idx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ══════════════════════════════════════════════════════════════════
# 6장. 하이퍼파라미터 튜닝
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '6. 하이퍼파라미터 튜닝', 1)
add_para(doc, '(Day 06에 해당)')
add_para(doc, 'Day 5에서 RF가 LightGBM보다 통계적으로 유의하게 우세함이 확인된 상태에서, LightGBM의 Calibration Gap이 크다는 점은 튜닝을 통해 개선 여지가 크다는 것을 시사함. LightGBM에 체계적인 하이퍼파라미터 탐색을 수행하여 이 가능성을 검증함.')

add_heading(doc, '6.1 LightGBM: Optuna 튜닝', 2)
add_para(doc, 'Optuna 프레임워크를 사용하여 베이지안 최적화 방식의 하이퍼파라미터 탐색을 수행함. 과적합 방지를 위해 n_estimators=5,000으로 고정하고 early_stopping_rounds=50을 적용하여 최적 트리 수를 자동으로 결정함.')

add_table(doc,
    ['설정 항목', '값'],
    [
        ('탐색 방법', 'Optuna (Tree-structured Parzen Estimator)'),
        ('총 trials', '20회'),
        ('n_estimators', '5,000 (고정) + early_stopping_rounds=50'),
        ('scale_pos_weight', '27.58 (고정, 클래스 불균형 비율)'),
        ('Pruner', 'MedianPruner(n_startup_trials=5, n_warmup_steps=2)'),
        ('총 소요 시간', '62분 47초'),
        ('Best trial', 'Trial 17'),
    ],
    col_widths=[5, 10])
add_caption(doc, '[표 6-1] LightGBM Optuna 튜닝 설정')

add_para(doc, 'Best 하이퍼파라미터 (Trial 17):', bold=True)
add_table(doc,
    ['파라미터', '탐색 결과값'],
    [
        ('learning_rate', '0.1563037874293778'),
        ('num_leaves', '165'),
        ('min_child_samples', '32'),
        ('subsample', '0.9248534606596772'),
        ('colsample_bytree', '0.6651160762107319'),
        ('reg_lambda', '0.00025077983643406634'),
    ],
    col_widths=[5, 10])
add_caption(doc, '[표 6-2] LightGBM Optuna Best 하이퍼파라미터')

add_para(doc, '튜닝 결과: PR-AUC 0.4860 → 0.5101 (+0.0241 개선). Day 5에서 통계적으로 유의하게 앞섰던 RF(0.4972)를 PR-AUC 기준으로 처음으로 역전함.')

add_heading(doc, '6.2 Random Forest: OOB 기반 탐색', 2)
add_para(doc, 'RF는 Optuna 탐색 시 trial당 25분 이상 소요되어 20 trials 기준 약 500분(8시간)이 필요함. 현실적인 실험 비용을 고려하여 Optuna 탐색을 보류하고, RF 고유의 OOB(Out-of-Bag) 스코어를 활용한 그리드 탐색으로 대체함.')

add_table(doc,
    ['max_depth', 'min_samples_leaf', 'max_features', 'OOB Score'],
    [
        ('None', '2', '0.5', '0.9812 (1위)'),
        ('None', '1', '0.3', '0.9803 (2위)'),
        ('20', '2', '0.3', '0.9722 (3위)'),
        ('15', '1', '0.5', '0.9552'),
        ('10', '1', '0.3', '0.9028'),
    ],
    col_widths=[2.5, 3.5, 3, 4])
add_caption(doc, '[표 6-3] RF OOB 기반 탐색 결과 (상위 5개)')

add_para(doc, '추가로 ExtraTreesClassifier를 비교한 결과 PR-AUC 0.4912로, RF 기본값(0.4972) 대비 -0.006을 기록하여 RF 기본값을 유지하기로 결정함.')

add_heading(doc, '6.3 튜닝 결과 및 최종 후보 선정', 2)
add_table(doc,
    ['모델', 'Day 5 PR-AUC', 'Day 6 PR-AUC', '변화'],
    [
        ('LightGBM (Optuna 튜닝)', '0.4860', '0.5101', '+0.0241 → RF 역전'),
        ('Random Forest (기본값 유지)', '0.4972', '0.4972', '0 (Optuna 비용 문제로 보류)'),
        ('XGBoost (비교 제외)', '0.4471', '—', '—'),
        ('Logistic Regression (베이스라인)', '0.3366', '—', '—'),
    ],
    col_widths=[5, 2.5, 2.5, 5])
add_caption(doc, '[표 6-4] 튜닝 전후 PR-AUC 변화 및 최종 후보 선정')

add_para(doc, '최종 후보: LightGBM(0.5101) vs RF(0.4972). PR-AUC 격차는 0.0129로 확대됨. 그러나 이 시점은 아직 Calibration 이전이므로, Day 7에서 확률 보정까지 완료한 후 최종 모델을 결정함.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7장. Calibration 및 Threshold Optimization
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '7. Calibration 및 Threshold Optimization', 1)
add_para(doc, '(Day 07에 해당)')

add_heading(doc, '7.1 확률 보정 결과 및 최종 모델 선택', 2)
add_para(doc, 'LightGBM과 RF 모두에 CalibratedClassifierCV(method=\'isotonic\', cv=5)를 적용하여 예측 확률을 실제 사기 비율에 맞게 보정함. Isotonic 보정은 단조 변환 제약이 없어 높은 보정 오차를 가진 모델에 더 효과적임.')

add_table(doc,
    ['단계', 'LightGBM PR-AUC', 'RF PR-AUC', '상황'],
    [
        ('Day 5 (기본 모델)', '0.4848', '0.4971', 'RF 통계적 우세 (CI: [0.0089, 0.0158])'),
        ('Day 6 (튜닝 후)', '0.5101', '0.4972', 'LightGBM 역전 시도'),
        ('Day 7 (보정 후)', '0.5517', '0.5100', 'LightGBM 최종 역전 확정'),
    ],
    col_widths=[4.5, 3, 2.5, 5])
add_caption(doc, '[표 7-1] Day 5→6→7 PR-AUC 변화 및 역전 과정')

add_table(doc,
    ['평가 기준', 'LightGBM (최종 채택)', 'RF'],
    [
        ('PR-AUC (mean±std)', '0.5517 ± 0.0272', '0.5100 ± 0.0339'),
        ('Calibration Gap (평균 |Gap|)', '0.1447 (보정 전 0.3130 대비 53.7% 개선)', '0.0507 (보정 전 0.1366 대비 62.9% 개선)'),
        ('Brier Score', '0.0241', '0.0231'),
        ('fold별 PR-AUC', '0.5070 ~ 0.5923', '0.4504 ~ 0.5561'),
    ],
    col_widths=[4.5, 6.5, 4])
add_caption(doc, '[표 7-2] 확률 보정 후 LightGBM vs RF 최종 비교')

add_para(doc, 'RF의 Calibration Gap(0.0507)과 Brier Score(0.0231)는 LightGBM보다 우수하나, 핵심 성능 지표인 PR-AUC에서 LightGBM(0.5517)이 RF(0.5100)를 0.0417 앞서므로 LightGBM을 최종 모델로 선택함.')
add_para(doc, '이로써 Day 5에서 통계적으로 유의하게 RF에 뒤처졌던 LightGBM이, 체계적인 하이퍼파라미터 튜닝(+0.0241)과 확률 보정(+0.0416)을 거쳐 PR-AUC 기준으로 RF를 최종 역전하는 과정이 완성됨.')

add_heading(doc, '7.2 Threshold Optimization (비용 최소화)', 2)
add_para(doc, '기본 threshold(0.5)는 불균형 데이터에서 사기 거래 대부분을 놓치는 결과를 초래함. 실무에서는 FN 비용(놓친 사기)과 FP 비용(오탐)의 비율에 따라 최적 threshold가 달라지므로, 이를 명시적으로 모델링함.')

add_para(doc, 'FN 비용 근거', bold=True)
add_table(doc,
    ['통계량', '값'],
    [
        ('사기 거래 건수', '20,663건'),
        ('TransactionAmt 중앙값', '75.00'),
        ('TransactionAmt 75분위수', '161.00'),
        ('TransactionAmt 평균', '149.24'),
    ],
    col_widths=[5, 5])
add_caption(doc, '[표 7-3] FN 비용 산정 근거 (사기 거래 TransactionAmt 분포)')

add_para(doc, '3가지 FN 기준(중앙값·75분위수·평균) × 6가지 FP 비용(1, 3, 5, 10, 20, 30) = 총 18개 시나리오를 분석하여 각 시나리오별 총 비용 최소화 threshold를 도출함.')
add_para(doc, '대표 시나리오 결과 (FN 기준: 중앙값 75.0):', bold=True)

add_table(doc,
    ['FP 비용', 'FN:FP 비율', '최적 threshold', 'Precision', 'Recall', 'F1', 'FN', 'FP'],
    [
        ('1', '75.0:1', '0.05', '0.0847', '0.8379', '0.1538', '3,349건', '187,155건'),
        ('3', '25.0:1', '0.12', '0.1883', '0.6967', '0.2964', '6,268건', '62,072건'),
        ('5', '15.0:1', '0.16', '0.2781', '0.6313', '0.3861', '7,619건', '33,863건'),
        ('10', '7.5:1', '0.20', '0.4008', '0.5608', '0.4675', '9,076건', '17,324건'),
        ('20', '3.8:1', '0.28', '0.6239', '0.4562', '0.5270', '11,237건', '5,682건'),
        ('30', '2.5:1', '0.33', '0.7477', '0.4114', '0.5307', '12,163건', '2,868건'),
    ],
    col_widths=[2, 2.5, 3, 2.5, 2, 2, 2.5, 3])
add_caption(doc, '[표 7-4] FN 기준: 중앙값(75.0) 시나리오별 최적 threshold 결과')

add_para(doc, '대표 시나리오 선택: FP 비용=10 (FN:FP=7.5:1) → threshold=0.20', bold=True)
add_table(doc,
    ['지표', '값'],
    [
        ('Threshold', '0.20'),
        ('Precision', '0.4008'),
        ('Recall', '0.5608'),
        ('F1-score', '0.4675'),
        ('TN (정상→정상)', '552,153건'),
        ('FP (정상→사기, 오탐)', '17,324건'),
        ('FN (사기→정상, 놓친 사기)', '9,076건'),
        ('TP (사기→사기, 탐지 성공)', '11,587건'),
    ],
    col_widths=[5, 5])
add_caption(doc, '[표 7-5] 최종 선택 threshold=0.20 성능 지표')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8장. 설명가능성 분석 (SHAP)
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '8. 설명가능성 분석 (SHAP)', 1)
add_para(doc, '(Day 08에 해당)')

add_heading(doc, '8.1 분석 모델 구성', 2)
add_para(doc, 'SHAP 분석은 목적에 따라 두 가지 모델을 명확히 구분하여 사용함.')

add_table(doc,
    ['구분', '파일명', '학습 방식', '타겟 인코딩', '사용 목적'],
    [
        ('운영용 모델', 'day8_calibrated_lgbm.pkl',
         'CalibratedClassifierCV\n(isotonic, cv=5)',
         'fold 내부 (누수 없음)',
         '실제 사기 탐지 배포용'),
        ('해석용 모델', 'day8_shap_model.pkl',
         '전체 데이터 재학습\n(best_iteration=642)',
         '전체 데이터 기준\n(의도적 적용)',
         'TreeSHAP 기반 전역 해석'),
    ],
    col_widths=[2.5, 4, 3.5, 3.5, 3.5])
add_caption(doc, '[표 8-1] 운영용 모델과 해석용 모델 구분')

add_para(doc, '해석용 모델에 전체 데이터 기준 타겟 인코딩을 적용한 것은 의도적 설계임. SHAP 전역 해석에서 모든 샘플에 대해 일관된 인코딩값을 부여하여 변수 중요도를 안정적으로 산출하기 위한 목적임. 이 설계는 시간적 Drift 검증과는 방법론적으로 구분되며, Day 9에서 해당 상충 관계를 실증적으로 재검증함.')

add_heading(doc, '8.2 전역 중요도 분석', 2)
add_para(doc, 'SHAP TreeExplainer를 사용하여 전역 중요도를 산출함. 분석 샘플: 사기 3,000건 + 정상 17,000건 = 20,000건 (np.random.seed=42).')

add_table(doc,
    ['순위', '변수', '해석'],
    [
        ('1위', 'card1', '특정 카드 번호 그룹이 사기 패턴의 핵심 신호. SHAP=1.492 (해석용 모델 기준)'),
        ('2위', 'C13', '값이 낮을수록 사기 확률 증가. 거래 카운팅 변수로 비정상 패턴 반영'),
        ('3위', 'C_PC1', 'C컬럼군 PCA 첫 번째 주성분'),
        ('4위', 'card2', '카드 번호 2 (card1과 함께 카드 식별 역할)'),
        ('5위', 'Transaction_day', '비선형 패턴. 특정 날짜에 사기 집중 가능성'),
        ('6위', 'P_emaildomain', '결제자 이메일 도메인. 특정 도메인이 고위험'),
        ('7위', 'C_PC2', 'C컬럼군 PCA 두 번째 주성분'),
        ('8위', 'dist1', '거래 장소와 주소 간 거리'),
        ('9위', 'addr1', '청구 주소 (지역 코드)'),
        ('10위', 'amt_robust_zscore', 'Day 3 설계 파생변수. 효과 검증됨'),
    ],
    col_widths=[1.5, 4, 9.5])
add_caption(doc, '[표 8-2] SHAP 전역 중요도 상위 10위 (day8_shap_model 기준)')

add_para(doc, 'Day 3에서 설계한 amt_robust_zscore가 SHAP 상위 10위에 진입하여, TransactionAmt 원본을 그대로 사용하는 것 대비 파생변수 설계의 실효성이 검증됨.')
add_para(doc, '※ card1의 SHAP 중요도(1위) 및 절대값의 타당성은 Day 9에서 시간적 누수 여부와 스무딩 과적합 여부를 통해 추가 검증됨. 결론적으로 card1의 1위 순위는 견고하게 재현되나, 절대적 SHAP 값은 저빈도 카테고리의 잔여 과적합으로 일부 과대추정됨이 확인됨.')

add_heading(doc, '8.3 개별 사례 해석', 2)
add_para(doc, '개별 거래에 대한 모델의 의사결정 근거를 Waterfall Plot으로 설명함. 대표 사례는 다음과 같음.')

add_table(doc,
    ['사례', 'f(x) (log-odds)', '주요 기여 변수', '해석'],
    [
        ('사기 거래 #5', '+5.013', 'card1(+2.83), C13(+2.02), C_PC2(+1.61)', '카드 그룹 특성과 C13 낮은 값이 함께 사기 방향으로 강하게 작용'),
        ('사기 거래 #6', '+4.574', 'C13(+3.66), ProductCD(+0.70), card1(+0.53)', 'C13이 압도적 기여. ProductCD 유형도 고위험 신호'),
        ('정상 거래 #0', '-10.672', 'card1(-1.48), card2(-0.82), C_PC1(-0.48)', '카드 그룹 특성이 정상 방향으로 강하게 억제'),
        ('정상 거래 #1', '-9.38', 'card1(+1.09)이나 V_group6(-0.82) 등이 억제', 'card1은 사기 방향이나 다수 변수가 상쇄하여 정상 판정'),
    ],
    col_widths=[2.5, 2.5, 5, 5])
add_caption(doc, '[표 8-3] Waterfall Plot 개별 사례 해석 요약')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9장. Drift 분석 및 방법론적 자기 검증
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '9. Drift 분석 및 방법론적 자기 검증', 1)
add_para(doc, '(Day 09에 해당 — 본 프로젝트의 방법론적 차별화가 가장 집약된 장)')

add_heading(doc, '9.1 분할 설계 및 목적 구분', 2)
add_para(doc, 'Day 04~08의 모델 학습·평가에 사용한 분할과 Day 09의 Drift 검증용 분할은 목적이 다른 별개의 분할임. 이 구분을 명확히 하는 것이 방법론적 일관성을 위해 중요함.')

add_table(doc,
    ['구분', '분할 기준', '목적', '적용 범위'],
    [
        ('Day 04~08 GroupKFold', 'UID(card1+card2+addr1) 기준\n(시간 순서 무관)',
         '모델 학습·평가용\n(데이터 누수 방지)',
         '성능 평가, 통계 검증,\n튜닝, Calibration, SHAP'),
        ('Day 09 시계열 분할', 'Transaction_month 0~4\nvs month 5+6',
         'Drift 검증 전용\n(시간적 분포 변화 시뮬레이션)',
         'PSI/KS 계산,\n보조 모델 재학습'),
    ],
    col_widths=[3.5, 4, 3.5, 4])
add_caption(doc, '[표 9-1] 두 가지 분할 방식의 목적 구분')

add_para(doc, '시계열 분할 상세:')
add_table(doc,
    ['구간', '거래 건수', '사기 비율', '비고'],
    [
        ('학습 구간 (month 0~4)', '495,904건', '3.50%', '—'),
        ('검증 구간 (month 5+6)', '94,636건', '3.47%', 'month 6은 2일치 8,111건 → month 5와 병합'),
    ],
    col_widths=[4, 3, 3, 5])
add_caption(doc, '[표 9-2] 시계열 분할 결과')

add_heading(doc, '9.2 연속형 변수 PSI / KS Test', 2)
add_para(doc, '연속형 변수 153개(Transaction_day 제외)에 대해 PSI(Population Stability Index)를 학습 구간 기준 percentile bin으로 계산하고, KS Test로 분포 차이의 통계적 유의성을 검증함.')
add_para(doc, '※ Transaction_day는 PSI=6.901로 시간 단조 증가에 의한 구조적 상승이므로 분석 대상에서 제외함.')

add_table(doc,
    ['PSI 등급', '기준', '변수 수'],
    [
        ('안정', 'PSI < 0.1', '143개'),
        ('주의', '0.1 ≤ PSI < 0.2', '4개'),
        ('Drift 위험', 'PSI ≥ 0.2', '6개'),
    ],
    col_widths=[3, 4, 3])
add_caption(doc, '[표 9-3] 연속형 변수 PSI 분류 결과 (153개 대상)')

add_table(doc,
    ['변수', 'PSI', 'KS 통계량', '해석'],
    [
        ('AllFlag_PC3', '0.3033', '0.1859', '결측 패턴 변화가 가장 큼'),
        ('AllFlag_PC1', '0.2911', '0.1804', '동상'),
        ('AllFlag_PC7', '0.2709', '—', 'AllFlag 계열 지속 Drift'),
        ('AllFlag_PC9', '0.2167', '—', '동상'),
        ('AllFlag_PC5', '0.2040', '—', '동상'),
        ('AllFlag_PC4', '0.2019', '—', '동상'),
        ('D11', '0.1629', '0.1538', '시간 간격 변수 분포 변화'),
    ],
    col_widths=[3.5, 2, 2.5, 7])
add_caption(doc, '[표 9-4] 연속형 변수 PSI 상위 결과 (Drift 위험 및 주의 등급)')

add_para(doc, 'KS Test 결과: 153개 중 130개(84.9%)에서 분포 차이가 유의함(p < 0.05). AllFlag(결측 패턴) 계열이 PSI와 KS 모두 상위를 차지하며, 시간에 따라 결측 패턴이 변화하고 있음을 시사함.')

add_heading(doc, '9.3 범주형 변수 PSI', 2)
add_para(doc, '범주형 변수 49개에 대해 카테고리별 비율 차이 기반 PSI를 계산함.')

add_table(doc,
    ['PSI 등급', '기준', '변수 수'],
    [
        ('안정', 'PSI < 0.1', '38개'),
        ('주의', '0.1 ≤ PSI < 0.2', '4개'),
        ('Drift 위험', 'PSI ≥ 0.2', '7개'),
    ],
    col_widths=[3, 4, 3])
add_caption(doc, '[표 9-5] 범주형 변수 PSI 분류 결과 (49개 대상)')

add_table(doc,
    ['변수', 'PSI (원본)', '카디널리티', '비고'],
    [
        ('card1', '2.4262', '고카디널리티 (학습: 12,890개)', '→ 9.4절 재검증'),
        ('id_31', '0.8319', '19개 (학습)', '→ 9.4절 재검증'),
        ('id_13', '0.7490', '20개 (학습)', '→ 9.4절 재검증'),
        ('M9', '0.2396', '3개', '확정 Drift'),
        ('M8', '0.2393', '3개', '확정 Drift'),
        ('M7', '0.2391', '3개', '확정 Drift'),
        ('M2', '0.2003', '3개', '확정 Drift'),
        ('M3', '0.1990', '3개', '확정 Drift'),
        ('M1', '0.1988', '3개', '확정 Drift'),
    ],
    col_widths=[2, 2.5, 4, 6.5])
add_caption(doc, '[표 9-6] 범주형 PSI 상위 결과 (Drift 위험 및 주의 등급)')

add_heading(doc, '9.4 고카디널리티 PSI 착시 발견 및 재검증', 2)
add_para(doc, 'card1, id_31, id_13 세 변수는 카디널리티가 매우 높아, 원본 PSI가 실제 분포 변화를 반영하는지 아니면 희소 카테고리의 수학적 착시인지 검증이 필요함.')
add_para(doc, '검증 방법: 상위 N개 카테고리 + Others 재그룹핑(regroup_top_n, top_n=20) 후 PSI 재계산.')

add_table(doc,
    ['변수', 'PSI 원본', 'PSI 재그룹핑 후', 'Other 비중(학습)', 'Other 비중(검증)', '판정'],
    [
        ('card1', '2.4262', '0.0029', '77.4%', '77.6%', 'Drift 위험 완전 제외'),
        ('id_31', '0.8319', '0.3498', '8.4%', '14.4%', '확정 Drift'),
        ('id_13', '0.7490', '0.5573', '1.0%', '4.4%', '확정 Drift'),
    ],
    col_widths=[2, 2, 2.5, 3, 3, 3.5])
add_caption(doc, '[표 9-7] 고카디널리티 변수 재그룹핑 PSI 비교')

add_para(doc, 'card1 재검증 결론: top20+Other 재그룹핑 후 PSI가 2.426에서 0.003으로 사실상 소멸했으며, Other 비중이 학습(77.4%)과 검증(77.6%)에서 거의 동일함. 즉 상위 카드 그룹의 점유율은 시간에 따라 변하지 않았고, 원본 PSI 2.426은 수천 개 희소 카테고리의 미세한 흔들림이 PSI의 로그 항에서 수학적으로 증폭된 구조적 착시였음이 실증됨. 이 발견은 고카디널리티 변수에 대해 원본 PSI를 그대로 해석할 때의 위험성을 보여줌.')
add_para(doc, 'id_31, id_13은 재그룹핑 후에도 PSI가 0.350, 0.557로 유지되어 확정 Drift로 판정함.')

add_heading(doc, '9.5 시간적 누수 발견 및 보조 모델 재검증', 2)
add_para(doc, 'Day 8의 해석용 모델(day8_shap_model)은 전체 데이터를 기반으로 타겟 인코딩을 재계산하여 학습됨. 이 모델로 시간적 Drift를 검증하면, 검증 구간 데이터의 정보가 학습 시 이미 인코딩에 반영되어 있는 문제가 발생함(시간적 누수, temporal leakage). 이를 실증 데이터로 확인함.')

add_para(doc, '시간적 누수 실증 (card1 카테고리별 인코딩값 비교):', bold=True)
add_table(doc,
    ['card1 카테고리', '학습구간 평균', '전체데이터 평균', '모델 인코딩값', '일치 여부'],
    [
        ('7919', '0.007709', '0.007501', '0.007537', '전체 평균과 일치'),
        ('9500', '0.035647', '0.037283', '0.037280', '전체 평균과 일치'),
        ('15885', '0.039480', '0.042853', '0.042838', '전체 평균과 일치'),
        ('17188', '0.029650', '0.026875', '0.026891', '전체 평균과 일치'),
        ('15066', '0.039394', '0.039396', '0.039385', '전체 평균과 일치'),
    ],
    col_widths=[3, 3, 3, 3, 3])
add_caption(doc, '[표 9-8] card1 타겟 인코딩값 비교 — 시간적 누수 실증')

add_para(doc, '모델 인코딩값이 학습 구간 평균이 아닌 전체 데이터 평균과 일치함이 확인됨. 이는 배포용 모델로서는 표준적이고 타당한 선택이나, 이 모델을 이용한 시간적 Drift 검증에는 그대로 사용할 수 없음을 의미함.')

add_para(doc, '보조 모델 재학습 및 재검증', bold=True)
add_para(doc, '학습 구간 데이터만으로 타겟 인코딩을 재계산(target_encode_smoothed, m=10)하고 동일 하이퍼파라미터로 보조 모델을 학습한 후, Score PSI를 재계산함.')

add_table(doc,
    ['모델', 'Score PSI', '학습 구간\n평균 예측확률', '검증 구간\n평균 예측확률', '판정'],
    [
        ('day8 최종모델 (누수 있음)', '0.0018', '0.1416', '0.1430', '안정'),
        ('보조모델 (누수 제거, m=10)', '0.0319', '0.0350', '0.0179', '안정 ※'),
    ],
    col_widths=[4.5, 2.5, 2.5, 2.5, 3])
add_caption(doc, '[표 9-9] 인코딩 방식에 따른 Score PSI 비교')

add_para(doc, '※ 보조 모델의 Score PSI는 0.0319로 여전히 안정(< 0.1) 등급이나, 누수 있는 모델 대비 약 17배 상승하여 더 현실적인 Drift 수준을 반영함. 더 중요한 신호는 검증 구간의 평균 예측 사기확률이 학습 구간(0.0350) 대비 약 50% 하락(0.0179)한다는 점임. 이는 PSI 단일 지표만으로는 포착되지 않는 형태의 잠재적 성능 저하 신호로, 모델이 검증 구간에서 사기 확률을 체계적으로 과소 추정하는 경향이 있음을 시사함.')

add_para(doc, 'SHAP 재계산 결과 (보조모델 m=10, 20,000건 샘플):', bold=True)
add_table(doc,
    ['순위', '변수', 'SHAP (보조모델 m=10)', 'SHAP (day8 누수 있음)'],
    [
        ('1위', 'card1', '2.589198', '1.492'),
        ('2위', 'C13', '1.888500', '—'),
        ('3위', 'C_PC1', '1.360016', '—'),
        ('4위', 'card2', '0.900344', '—'),
        ('5위', 'Transaction_day', '0.677267', '—'),
    ],
    col_widths=[1.5, 4, 4.5, 4.5])
add_caption(doc, '[표 9-10] SHAP 상위 5위 — 인코딩 방식 비교 (재계산 결과 상위 20위 기준)')

add_para(doc, 'card1의 SHAP 1위 순위는 인코딩 방식과 무관하게 안정적으로 재현됨. 또한 Drift 위험 변수 14개(AllFlag_PC 계열, M1~M3·M7~M9, id_31, id_13, D11 등) 중 어느 것도 두 모델 모두에서 SHAP 상위 20위 안에 진입하지 않음. 이는 모델이 소수의 불안정한 변수에 과도하게 의존하지 않는다는 것을 인코딩 방식과 무관하게 일관적으로 확인하는 결과임.')

add_heading(doc, '9.6 card1 스무딩 강도 비교', 2)
add_para(doc, 'card1은 SHAP 1위임과 동시에 학습 빈도가 낮은 카테고리(저빈도 카테고리)가 수천 개 포함됨. 저빈도 카테고리에서 타겟 인코딩의 개별 샘플 수준 과적합이 SHAP 값에 영향을 미치는지 스무딩 강도(m=10 vs m=100)를 비교하여 검증함.')
add_para(doc, '스무딩 공식: smoothed_mean = (n × category_mean + m × global_mean) / (n + m). m이 클수록 저빈도 카테고리가 전체 평균으로 수렴함.')

add_table(doc,
    ['빈도 구간', 'm=10 SHAP', 'm=100 SHAP', '해석'],
    [
        ('1~5건', '1.389', '2.195', '—'),
        ('6~50건', '2.434', '2.032', '—'),
        ('51~500건 ※', '3.799', '2.060', '과적합 해소 (1.739 감소)'),
        ('501~5,000건', '2.403', '2.328', '거의 동일'),
        ('5,000건 초과 ※', '1.556', '1.550', '안정 (진짜 신호 확인)'),
    ],
    col_widths=[3, 2.5, 2.5, 7])
add_caption(doc, '[표 9-11] card1 빈도 구간별 평균 |SHAP| — 스무딩 강도 비교')

add_table(doc,
    ['모델', 'card1 SHAP', '순위'],
    [
        ('day8 최종모델 (누수 있음)', '1.492', '1위'),
        ('보조모델 m=10 (학습구간 전용 인코딩)', '2.589', '1위'),
        ('보조모델 m=100 (강한 스무딩)', '2.131', '1위'),
    ],
    col_widths=[7, 3, 2])
add_caption(doc, '[표 9-12] card1 SHAP 값 — 인코딩 방식 및 스무딩 강도별 비교')

add_para(doc, 'm=10에서 관찰된 51~500건 구간의 이상 급등(3.799)은 스무딩 강화로 2.060으로 감소하여 해당 구간의 SHAP 상승이 과적합에 기인한 것임이 확인됨. 반면 표본이 충분한 5,000건 초과 구간의 SHAP는 스무딩 강도와 무관하게 일정하게 유지되어(1.556→1.550), card1이 실제로 카드 그룹별 사기율 차이를 반영하는 유의미한 예측 신호를 담고 있음을 뒷받침함.')
add_para(doc, '최종 판단: card1의 SHAP 최상위권(1위)은 인코딩 방식·스무딩 강도와 무관하게 타당함. 단, 절대적 SHAP 값(m=10: 2.589)은 저빈도 카테고리의 잔여 과적합에 의해 다소 과대추정된 수치이므로, card1 중요도 해석 시 순위·존재의 의미에 무게를 두고 절대값 해석에는 주의가 필요함.')

add_heading(doc, '9.7 종합 결론', 2)

add_table(doc,
    ['결론', '내용'],
    [
        ('1. 입력 변수 수준 Drift',
         '결측 패턴(AllFlag_PC 계열), 검증 로직(M1~M3·M7~M9), 기기 정보(id_13·id_31)에서 시간에 따른 실질적 분포 변화가 확인됨'),
        ('2. 모델 의존도 안정성',
         'Drift 위험 변수 14개 모두 SHAP 상위 20위 밖. 인코딩 방식과 무관하게 재현됨 → 모델이 불안정 변수에 과의존하지 않음'),
        ('3. Score PSI 한계',
         '검증 구간 평균 예측 사기확률이 0.0350→0.0179로 50% 하락하는 것은 PSI 단일 지표로 포착 불가. 분포의 위치 변화를 함께 모니터링해야 함'),
        ('4. 배포 모델과 검증 모델의 분리 필요성',
         'day8 최종 모델(전체데이터 학습)은 배포용으로 타당. 단, 시간적 Drift 검증에는 반드시 학습구간 전용 재인코딩 모델을 사용해야 함'),
        ('5. 고카디널리티 타겟 인코딩 주의사항',
         'card1과 같은 고카디널리티 식별자 변수는 시간적 누수뿐 아니라 스무딩 강도에 따른 개별 샘플 과적합 여부도 별도로 점검 필요'),
    ],
    col_widths=[3.5, 11.5])
add_caption(doc, '[표 9-13] Day 9 종합 결론 5가지')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10장. 종합 결론 및 실무적 시사점
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '10. 종합 결론 및 실무적 시사점', 1)

add_heading(doc, '10.1 프로젝트 성과 요약', 2)
add_para(doc, '본 프로젝트는 IEEE-CIS Fraud Detection 데이터셋을 활용하여, Day 01부터 Day 09까지 9단계의 체계적인 분석 과정을 거쳐 비용 민감형 사기탐지 모델을 완성함.')

add_table(doc,
    ['단계', '주요 성과'],
    [
        ('데이터 탐색 (Day 01~02)',
         '클래스 불균형(1:27.6) 정량화, V컬럼 15개 구조적 그룹 발견, 통계 검정으로 유의 변수 선별'),
        ('Feature Engineering (Day 03)',
         'UID 기반 사용자 맥락 변수 설계, amt_robust_zscore(effect size +16.6% 개선), V컬럼 339→89개 PCA'),
        ('모델 비교 (Day 04)',
         'StratifiedGroupKFold 교체로 평가 지표 왜곡 수정, 4개 모델 OOF 비교 완료'),
        ('통계적 검증 (Day 05)',
         '3종 검정(DeLong·Bootstrap CI·McNemar) + Holm 보정으로 모델 순위 통계적 확정'),
        ('튜닝 + 보정 (Day 06~07)',
         'LightGBM Optuna 튜닝으로 PR-AUC 0.4860→0.5101, Calibration 후 0.5517로 최종 역전'),
        ('Threshold 최적화 (Day 07)',
         'FN:FP=7.5:1 비용 구조에서 threshold=0.20 확정 (Recall=0.5608, Precision=0.4008)'),
        ('설명가능성 (Day 08)',
         'card1·C13·C_PC1 등 상위 변수 SHAP 해석, 파생변수 효과 검증'),
        ('Drift 분석 (Day 09)',
         '고카디널리티 PSI 착시 실증, 시간적 누수 발견 및 보조 모델 재검증, 스무딩 과적합 진단'),
    ],
    col_widths=[4, 11])
add_caption(doc, '[표 10-1] 단계별 주요 성과 요약')

add_heading(doc, '10.2 최종 모델 성능', 2)
add_table(doc,
    ['지표', '값', '비고'],
    [
        ('모델', 'LightGBM + CalibratedClassifierCV(isotonic)', 'day8_calibrated_lgbm.pkl'),
        ('PR-AUC', '0.5517 ± 0.0272', 'OOF 5-fold 기준'),
        ('ROC-AUC', '0.8463 (Day 4 기준)', '—'),
        ('Calibration Gap', '0.1447', '보정 전 0.3130 대비 53.7% 개선'),
        ('Brier Score', '0.0241', '—'),
        ('최적 Threshold', '0.20', 'FN:FP=7.5:1 비용 구조'),
        ('Threshold 기준 Precision', '0.4008', '—'),
        ('Threshold 기준 Recall', '0.5608', '—'),
        ('Threshold 기준 F1', '0.4675', '—'),
    ],
    col_widths=[5, 5, 5])
add_caption(doc, '[표 10-2] 최종 모델 성능 요약')

add_heading(doc, '10.3 실무적 시사점', 2)

add_para(doc, '① 평가 지표 선택의 중요성', bold=True)
add_para(doc, '클래스 불균형(1:27.6)이 심각한 사기탐지 환경에서는 ROC-AUC보다 PR-AUC가 실질적 탐지 성능을 더 직접적으로 반영함. 또한 예측 확률의 신뢰도(Calibration)가 Threshold 최적화의 품질을 결정하므로, 성능 지표와 Calibration 지표를 함께 모니터링해야 함.')

add_para(doc, '② 비용 구조 반영의 실무적 가치', bold=True)
add_para(doc, 'Threshold를 0.5 기본값 대신 비용 구조(FN:FP=7.5:1)에 맞게 0.20으로 조정함으로써 Recall이 56.1%로 유지됨. 비용 비율이 변화하면 threshold를 재최적화하는 것만으로 대응 가능한 유연한 구조를 확보함.')

add_para(doc, '③ 통계적 모델 비교의 필요성', bold=True)
add_para(doc, 'Day 5에서 RF가 통계적으로 유의하게 우세했으나, 튜닝과 보정을 거쳐 LightGBM이 역전되었음. 이는 단일 시점의 성능 수치로 모델을 조기에 확정하는 것의 위험성을 보여주며, 개선 여지(Calibration Gap, 튜닝 여력)까지 고려한 종합적 판단이 필요함을 시사함.')

add_para(doc, '④ 배포 모델과 검증 모델의 분리 원칙', bold=True)
add_para(doc, 'Day 9에서 발견한 핵심 교훈: 전체 데이터로 학습한 배포 모델과 시간적 Drift 검증에 사용하는 모델은 반드시 분리해야 함. 배포 모델은 가용 데이터 전체로 최적 학습한 것이 맞으나, 이를 그대로 Drift 모니터링에 사용하면 검증 구간 정보가 인코딩에 이미 반영되어 있어 왜곡된 결과를 초래함.')

add_para(doc, '⑤ PSI 단일 지표의 한계', bold=True)
add_para(doc, 'Score PSI가 0.0319로 안정 등급이더라도, 검증 구간의 평균 예측 사기확률이 0.0350에서 0.0179로 50% 하락하는 것은 PSI가 포착하지 못하는 잠재적 성능 저하 신호임. 분포의 위치(평균) 변화를 함께 추적하는 모니터링 체계가 필요함.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11장. 한계 및 향후 과제
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '11. 한계 및 향후 과제', 1)

add_heading(doc, '11.1 현재 분석의 한계', 2)

add_table(doc,
    ['한계 사항', '내용'],
    [
        ('단일 데이터셋 한정',
         'IEEE-CIS 데이터셋 내에서의 시간적 안정성만 점검함. 실제 운영 환경의 새로운 사기 패턴이나 거래 규모 변화에 대한 외부 타당성은 별도 검증이 필요함'),
        ('Drift 검증의 시뮬레이션 성격',
         'Day 9의 시계열 분할은 동일 데이터셋 내 시간적 분포 안정성 점검 수준임. 실제 미래 데이터 유입 시의 Drift와는 구분해야 함'),
        ('V컬럼 해석의 한계',
         'V컬럼 339개가 익명화되어 있어 PCA 주성분의 비즈니스 의미를 직접 해석할 수 없음'),
        ('RF 하이퍼파라미터 탐색 미완',
         'LightGBM은 Optuna 20 trials를 완료했으나 RF는 컴퓨팅 비용 문제로 OOB 기반 탐색에 그침. RF의 잠재적 최적 성능이 확인되지 않음'),
        ('UID의 불완전성',
         'card1+card2+addr1 기반 UID가 완벽한 개인 식별자가 아니므로, 사용자 맥락 변수의 신뢰도에 한계가 있음 (Fold 4 동시 하락 현상과 관련)'),
    ],
    col_widths=[4, 11])
add_caption(doc, '[표 11-1] 현재 분석의 주요 한계')

add_heading(doc, '11.2 2단계 확장 방향 — MLOps 파이프라인화', 2)
add_para(doc, '1단계에서 확립한 분석 로직(전처리, Feature Engineering, 모델 학습·평가, Threshold 최적화, Drift 분석)을 자동 실행 가능한 파이프라인 컴포넌트로 리팩토링하는 것이 2단계의 핵심 원칙임. 1단계의 통계적 완성도를 훼손하지 않고 그 위에 자동화 계층을 얹는 것이 목표임.')

add_table(doc,
    ['영역', '도구(후보)', '1단계 연계 내용'],
    [
        ('워크플로 오케스트레이션',
         'Airflow 또는 Prefect',
         '데이터 전처리·Feature Engineering(Day 01~03)을 ETL 파이프라인으로 자동화'),
        ('실험/모델 관리',
         'MLflow',
         '모델 비교·통계검증·Threshold 결과(Day 04~07)를 실험 단위로 버전 관리'),
        ('재학습 자동화',
         'Airflow + MLflow 연계',
         '신규 데이터 유입 시 모델 학습~Threshold 최적화 절차를 자동 재실행'),
        ('Drift 모니터링 자동화',
         'MLflow + 커스텀 알림',
         'Day 09의 PSI/KS Test를 주기적 배치 작업으로 전환, 임계치 초과 시 알림'),
    ],
    col_widths=[3.5, 3.5, 8])
add_caption(doc, '[표 11-2] 2단계 확장 방향 (MLOps 파이프라인)')

add_para(doc, '2단계 설계 원칙 — Day 9 교훈의 반영', bold=True)
add_para(doc, 'Day 9에서 발견한 "배포 모델과 시간적 검증용 모델이 분리되어야 한다"는 교훈은 2단계 파이프라인 설계에 다음과 같이 반영되어야 함. 재학습 파이프라인은 신규 데이터가 유입될 때마다 가용한 전체 데이터로 배포 모델을 재학습하는 것이 올바른 방식임. 반면 Drift 모니터링은 재학습 직전 구간의 데이터를 시간 기준 검증 데이터셋으로 별도 유지하고, 이 데이터셋에서 학습구간 전용 재인코딩을 수행한 모델로 PSI/KS를 계산해야 함. 배포 모델과 모니터링 모델을 동일 아티팩트로 공용하는 설계는 Day 9에서 실증된 누수 문제를 자동화 환경에서 반복하는 구조적 오류임.')

add_para(doc, '2단계의 세부 아키텍처(파이프라인 DAG 설계, 모델 레지스트리 구조 등)는 1단계 완료 후 별도 기획안으로 작성 예정임.')

add_heading(doc, '11.3 향후 개선 과제', 2)
add_table(doc,
    ['과제', '내용'],
    [
        ('RF 최적 성능 확인',
         'GPU 또는 클라우드 컴퓨팅 환경에서 RF Optuna 탐색 수행. LightGBM과의 최종 성능 비교 완성'),
        ('대표변수 버전 재평가',
         'Day 3에서 결측 처리를 보류한 대표변수 버전(129개 변수)의 처리 후 성능을 PCA 버전과 비교'),
        ('SMOTE 보완 실험',
         'V컬럼 익명성 문제가 해소되는 경우(또는 V컬럼 제외 버전에서) SMOTE와 scale_pos_weight의 성능 차이를 실험적으로 검증'),
        ('card1 인코딩 전략 고도화',
         '스무딩 m=100이 적정한지 교차검증 기반으로 최적 m값을 탐색하거나, 빈도 구간별 차등 스무딩 적용 방안 검토'),
        ('Drift 자동 알림 임계값 설정',
         'PSI 기준(0.1/0.2)과 함께 평균 예측확률 변화율(예: ±30% 이상 시 경보)을 복합 조건으로 설정하는 모니터링 정책 수립'),
    ],
    col_widths=[4, 11])
add_caption(doc, '[표 11-3] 향후 개선 과제')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 최종 저장
# ══════════════════════════════════════════════════════════════════
final_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                          '비용민감형_사기탐지_최종보고서.docx')
doc.save(final_path)
print(f"최종 보고서 저장 완료: {final_path}")
